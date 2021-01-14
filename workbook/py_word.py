import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION


'''-------------------------------- 读取excel表数据整理后写入word （纯文本）-----------------------------------'''
def wu_to_word(filepath):
    df = pd.read_excel(filepath, sheet_name="无")
    date_list = list(df['日期'])
    for d in date_list:
        filename = wordname+str(d)+"）.docx"  # 输出的word文件名
        title = "("+str(d)[:4]+"."+str(d)[4:6]+"."+str(d)[6:8]+")"  # 副标题日期XXXX.XX.XX
        word = str(d)[:4]+"年"+str(d)[4:6]+"月"+str(d)[6:8]+"日"  # 开头、落款日期XXXX年XX月XX日
        wu_doc(title, word, filename)
        print(f"文件：{filename},{title},{word} 已保存")

def wu_doc(title,word,filename):
    doc = Document()
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    # 段落的全局设置
    doc.styles['Normal'].font.name = u'宋体'  # 字体
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 中文字体需再添加这个设置
    doc.styles['Normal'].font.size = Pt(14)  # 字号 四号对应14
    t1 = doc.add_paragraph()  # 大标题
    t1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
    _t1 = t1.add_run(all_title)
    _t1.bold = True
    _t1.font.size = Pt(22)
    t2 = doc.add_paragraph()  # 副标题
    t2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
    _t2 = t2.add_run(title + "\n")
    _t2.bold = True
    doc.add_paragraph(word + "无单位报送数据。\n\n").paragraph_format.first_line_indent = Inches(0.35)  # 首行缩进
    doc.add_paragraph(word).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 右对齐
    doc.save(dir+filename)


'''-------------------------------- 读取excel表数据整理后写入word（需写入表格） -----------------------------------'''
def what_to_word(filepath):
    df = pd.read_excel(filepath, sheet_name="有")
    df.fillna('', inplace=True)  # 替换nan值为空字符
    dates = []  # 日期列表
    df_total = []  # 分日期存的所有df
    list_total = []  # 每一份word中需要的表数据合集
    for d in df.groupby('日期'):
        dates.append(d[0])
        df_total.append(d[1])
    for index,date in enumerate(dates):
        list_oneday = []  # 某一个word所需的表数据
        for row in range(len(df_total[index])):
            list_row = get_table_data(df_total, index, row)  # 其中一行数据
            list_oneday.append(list_row)
        list_total.append(list_oneday)
    for index, date in enumerate(dates):
        filename = wordname+str(date)+"）.docx"  # 输出的word文件名
        title = "("+str(date)[:4]+"."+str(date)[4:6]+"."+str(date)[6:8]+")"  # 副标题日期XXXX.XX.XX
        word = str(date)[:4]+"年"+str(date)[4:6]+"月"+str(date)[6:8]+"日"  # 开头、落款日期XXXX年XX月XX日
        sentence = get_sentence(df_total, index)
        what_doc(title, word, sentence, list_total[index], filename)
        print(f"文件：{filename} 已保存")
        # if index == 0:
        #     break


def get_table_data(df_total, df_index, table_row):
    list1 = df_total[df_index].values[table_row]  # excel表中的一行
    list2 = list1[3:7]  # 一至四级指标
    for i in range(len(list2)):  # 当前指标为空则沿用上级指标
        if list2[i] == '空' and i != 0:
            list2[i] = list2[i - 1]
    content = list1[2] + ":\n" + list1[-4]  # 报送内容
    if '否' in list1[-2]:  # 备注
        remark = '有记录未上传，' + str(list1[-1])
    else:
        remark = '已上传'
    list3 = list2.tolist()  # 需填入word中的表数据，由numpy数组转为list列表
    list3.append(str(content))
    list3.append(str(remark))
    return list3

def get_sentence(df_total, df_index):
    df_oneday = df_total[df_index]
    num = df_oneday['填报部门'].nunique()  # 报送单位的数量
    group = []  # 部门名称
    detail = []  # 组合某个部门的数据，其中元素为元组格式(, , , )
    info = ''  # 报送情况描述
    for item in df_oneday.groupby('填报部门'):
        group.append(item[0])
        detail.append(
            list(
                zip(
                    list(item[1]['报送内容']),
                    list(item[1]['记录数']),
                    list(item[1]['是否上报']),
                    list(item[1]['备注'])
                )
            )
        )
    for index, g in enumerate(group):  # 整理每个部门的填报情况
        mes = str(g)+'：'  # 部门开头
        for i in range(len(detail[index])):
            _mes = detail[index][i]
            if int(_mes[1])>0:
                mes = mes + f'“{_mes[0]}”{_mes[1]}条记录；'
        info = info + mes
    info = info[:-1]+"。"
    sentence = f"有{num}个部门报送了数据：{info}"
    return sentence

def what_doc(title, word, sentence, table, filename):  # 传入副标题日期，开头/落款日期，文段，表数据，文件名
    doc = Document()
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    # 段落的全局设置
    doc.styles['Normal'].font.name = u'宋体'  # 字体
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 中文字体需再添加这个设置
    doc.styles['Normal'].font.size = Pt(14)  # 字号 四号对应14
    t1 = doc.add_paragraph()  # 大标题
    t1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
    _t1 = t1.add_run(all_title)
    _t1.bold = True
    _t1.font.size = Pt(22)
    t2 = doc.add_paragraph()  # 副标题
    t2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
    _t2 = t2.add_run(title + "\n")
    _t2.bold = True
    doc.add_paragraph(word + sentence +"\n\n").paragraph_format.first_line_indent = Inches(0.35)  # 首行缩进
    doc.add_paragraph(word).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 右对齐
    doc.add_paragraph("各部门具体报送情况见附件：")

    doc.add_page_break()  # 分页---------------------------------------------------------------
    fujian = doc.add_paragraph().add_run("\n附件")
    fujian.bold = True
    fujian.font.size = Pt(16)
    t3 = doc.add_paragraph()  # 附件大标题
    t3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
    _t3 = t3.add_run("XX公司业务数据表")
    _t3.bold = True
    _t3.font.size = Pt(22)

    rows = len(table)+1
    word_table = doc.add_table(rows=rows, cols=6, style='Table Grid')  # 创建表格-------------
    word_table.autofit=True
    table = [table_title] + table
    for row in range(rows):  # 写入表格
        cells = word_table.rows[row].cells
        for col in range(6):
            cells[col].text = str(table[row][col])
    for i in range(len(word_table.rows)):    # 遍历行列，逐格修改样式
        for j in range(len(word_table.columns)):
            for par in word_table.cell(i, j).paragraphs:  # 修改字号
                for run in par.runs:
                    run.font.size = Pt(10.5)
            for par in word_table.cell(0, j).paragraphs:  # 第一行加粗
                for run in par.runs:
                    run.bold = True
    doc.save(dir+filename)



dir = ""
workbook = "workbook.xlsx"
wordname = "XX公司业务数据表（日报"
all_title = "XX公司业务报告"
table_title = ['一级指标', '二级指标', '三级指标', '四级指标', '各部门报送情况', '备注']

wu_to_word(dir+workbook)
what_to_word(dir+workbook)